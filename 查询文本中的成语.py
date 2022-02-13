import xlrd
from docx import Document
import re
import matplotlib.pyplot as plt
import numpy as np
from openpyxl import Workbook
import sqlite3
import pandas

#################先把成语从数据库中拿出来
#
# conn = sqlite3.connect("chinese-idioms-12976.db")
# cursor = conn.cursor()
# sql ="""
# select char1,char2,char3,char4 from main.idiom
# """
# cursor.execute(sql)
# result = cursor.fetchall()
# conn.close()
# mybook = Workbook()
# wa =mybook.active
# for i in range(len(result)):
#     wa.append([result[i][0]+result[i][1]+result[i][2]+result[i][3]])
# mybook.save("成语.xls")
# # print(result[0][0]+result[0][1]+result[0][2]+result[0][3])
#
import re
################### 在excel 中查询
file = pandas.read_excel(r"./成语.xlsx", sheet_name="1")
document = Document(r"翻译-表格(3).docx")
a = []
for j in range(len(document.paragraphs)):
    if document.paragraphs[j].text==[] or len(document.paragraphs[j].text)<=8 or  "2018年" in document.paragraphs[j].text:
        if 0<len(document.paragraphs[j].text)<=8:
            print("标题:",document.paragraphs[j].text)
        elif "2018年" in document.paragraphs[j].text:
            print("题注",document.paragraphs[j].text)

    elif document.paragraphs[j].text!=[]:
            for i in range(len(file)):
                data = file.loc[i][0]
                if data in document.paragraphs[j].text:
                    a.append(data)
                    print("第%d个段落" % j)

# for i in file:
#     print(i)
print(a)
['自得其乐', '相得益彰', '璀璨夺目', '黄金时代', '不得而知', '价值连城', '一模一样', '从容应对', '接二连三', '锦上添花', '半信半疑', '语重心长', '深恶痛绝', '望而却步', '忠心耿耿', '摇身一变', '通力合作', '就事论事', '不折不扣', '大街小巷', '总而言之', '容光焕发', '淋漓尽致', '自告奋勇', '自力更生', '相映成趣', '难以置信', '赞不绝口', '浑然一体', '风靡一时', '神气活现', '井井有条', '重男轻女', '别有风味', '点睛之笔', '浑然一体', '引人注目', '街头巷尾', '屡见不鲜', '聚精会神', '字里行间', '无伤大雅', '坚忍不拔', '焕然一新', '戮力同心', '势不可挡']


